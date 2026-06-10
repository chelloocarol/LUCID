"""
Apply Tier 1 text-only changes to a copy of the paper docx.
Creates 面向煤矿..._修订版T1.docx (does NOT overwrite original).

Changes applied:
  T1-2: Add FLOPs measurement resolution note to Table 1 paragraph (P79)
  T1-3: Add dataset access information to §3.1 (P69)

T1-1 (GARC formula) is NOT modified here because:
  - The formula is an OMML equation object; python-docx cannot read/write equation XML reliably
  - A separate comment file (GARC_formula_correction.md) should be provided instead

Usage:
    python tools/apply_tier1_changes.py
"""

import shutil
import os
from docx import Document

ORIG = "面向煤矿井下图像的可见度条件自适应与眩光校准复原方法_新颖性增强版.docx"
COPY = "面向煤矿井下图像的可见度条件自适应与眩光校准复原方法_修订版T1.docx"

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

ORIG_PATH = os.path.join(ROOT, ORIG)
COPY_PATH = os.path.join(ROOT, COPY)


def apply_t1_2(doc):
    """T1-2: Add FLOPs resolution note to the paragraph that mentions FLOPs (P79)."""
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if 'FLOPs由48.7 G增至49.6 G' in t:
            # Find the run that contains "FLOPs由48.7 G增至49.6 G"
            for run in para.runs:
                if 'FLOPs由48.7 G增至49.6 G' in run.text:
                    # Add note after the FLOPs reference
                    run.text = run.text.replace(
                        'FLOPs由48.7 G增至49.6 G',
                        'FLOPs由48.7 G增至49.6 G（FLOPs均在[XXX×YYY]分辨率单帧输入下用thop计算，作者需补充具体分辨率）'
                    )
                    print(f"  T1-2 applied to P{i}")
                    return True
    # If runs don't contain the text (mixed formatting), modify paragraph text directly
    for i, para in enumerate(doc.paragraphs):
        full_text = para.text
        if 'FLOPs由48.7 G增至49.6 G' in full_text:
            # Para has complex formatting, add a sentence at the end
            if para.runs:
                last_run = para.runs[-1]
                suffix = '（注：表1中FLOPs数值在[作者需补充：XXX×YYY分辨率]单帧输入下用thop计算，以确保与实测1080p FLOPs 88.5G/88.7G的差异可追溯。）'
                last_run.text = last_run.text + suffix
                print(f"  T1-2 fallback applied to P{i} (last run)")
                return True
    print("  WARNING: T1-2 target paragraph not found")
    return False


def apply_t1_3(doc):
    """T1-3: Add dataset access info to §3.1 paragraph (P69)."""
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if '本文构建了一个真实井下煤矿图像数据集' in t and '1080p固定监控摄像机' in t:
            # Add dataset access sentence at end of paragraph
            if para.runs:
                last_run = para.runs[-1]
                access_note = '数据集代理版本（40视频，n=152测试对，RIDCP伪标签参考）已公开于 github.com/chelloocarol/lucidmine-40-video-dataset；原始1920对数据集因商业数据保密协议暂不公开，如需访问请联系通讯作者。'
                # Add note on a new sentence
                if not last_run.text.endswith('。') and not last_run.text.endswith('）'):
                    last_run.text = last_run.text + '。'
                last_run.text = last_run.text + access_note
                print(f"  T1-3 applied to P{i}")
                return True
    print("  WARNING: T1-3 target paragraph not found")
    return False


def main():
    if not os.path.exists(ORIG_PATH):
        print(f"ERROR: {ORIG_PATH} not found")
        return

    # Copy original
    shutil.copy2(ORIG_PATH, COPY_PATH)
    print(f"Created copy: {COPY_PATH}")

    # Load copy
    doc = Document(COPY_PATH)
    print(f"Loaded {len(doc.paragraphs)} paragraphs")

    # Apply T1-2
    print("\nApplying T1-2 (FLOPs resolution note)...")
    apply_t1_2(doc)

    # Apply T1-3
    print("Applying T1-3 (dataset access info)...")
    apply_t1_3(doc)

    # Save
    doc.save(COPY_PATH)
    print(f"\nSaved: {COPY_PATH}")
    print("\nNOTE: T1-1 (GARC formula) NOT applied - see AUDIT_FINAL_SUMMARY.md for manual instruction")
    print("The correct GARC formula should be:")
    print("  M(x) = (1 − G(x)) · (0.5 + 0.5·V(x))")
    print("  S_final = clip(1 + γ_G · S · M(x), 0.55, 1.45)")


if __name__ == "__main__":
    main()
